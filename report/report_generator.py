import os
from datetime import datetime
from jinja2 import Template


class ReportGenerator:
    """Genereert HTML, DOCX en PDF rapporten vanuit rapport-JSON."""

    def generate_html(self, report_data: dict, output_path: str) -> None:
        template_path = os.path.join(
            os.path.dirname(__file__), '..', 'templates', 'report_template.html'
        )
        with open(template_path, 'r', encoding='utf-8') as f:
            template_str = f.read()

        html = Template(template_str).render(
            report=report_data,
            generated_at=datetime.now().strftime('%d-%m-%Y %H:%M'),
        )

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)

    def generate_docx(self, report_data: dict, output_path: str) -> None:
        from docx import Document
        from docx.shared import Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        for section in doc.sections:
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)

        meta = report_data.get('metadata', {})

        # ── Titel ──────────────────────────────────────────────────────── #
        title = doc.add_heading('Audit Bevindingenoverzicht', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub = doc.add_paragraph(
            f"Gegenereerd: {datetime.now().strftime('%d-%m-%Y')}  |  "
            f"Kijktermijn: {meta.get('kijktermijn_jaren', 'N/A')} jaar  |  "
            f"Bestanden geanalyseerd: {meta.get('bestanden_geanalyseerd', 'N/A')}"
        )
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        doc.add_page_break()

        # ── Samenvatting ───────────────────────────────────────────────── #
        samenvatting = report_data.get('samenvatting', '')
        if samenvatting:
            doc.add_heading('Samenvatting', 1)
            doc.add_paragraph(samenvatting)

        # ── Bevindingen ────────────────────────────────────────────────── #
        bevindingen = report_data.get('bevindingen', [])
        if bevindingen:
            doc.add_heading('Relevante Bevindingen', 1)

            for prioriteit in ['Hoog', 'Middel', 'Laag']:
                groep = [b for b in bevindingen if b.get('prioriteit') == prioriteit]
                if not groep:
                    continue

                doc.add_heading(f'Prioriteit: {prioriteit}', 2)

                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                for i, h in enumerate(['Bron', 'Jaar', 'Bevinding', 'Relevantie']):
                    hdr[i].text = h
                    hdr[i].paragraphs[0].runs[0].bold = True

                for b in groep:
                    row = table.add_row().cells
                    row[0].text = str(b.get('bron', ''))
                    row[1].text = str(b.get('jaar', ''))
                    row[2].text = str(b.get('bevinding', ''))
                    row[3].text = str(b.get('relevantie', ''))

                doc.add_paragraph()

        # ── Aandachtspunten ────────────────────────────────────────────── #
        aandachtspunten = report_data.get('aandachtspunten', [])
        if aandachtspunten:
            doc.add_heading('Aandachtspunten voor de Auditor', 1)
            for ap in aandachtspunten:
                p = doc.add_paragraph()
                p.add_run(ap.get('onderwerp', '')).bold = True
                doc.add_paragraph(ap.get('toelichting', ''))

        doc.save(output_path)

    def generate_pdf(self, report_data: dict, output_path: str) -> None:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        )

        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            leftMargin=2.5 * cm,
            rightMargin=2.5 * cm,
            topMargin=2.5 * cm,
            bottomMargin=2.5 * cm,
        )

        styles = getSampleStyleSheet()
        style_title = ParagraphStyle('title', parent=styles['Title'], fontSize=18, spaceAfter=6)
        style_sub = ParagraphStyle('sub', parent=styles['Normal'], fontSize=9,
                                   textColor=colors.HexColor('#64748B'), spaceAfter=20)
        style_h1 = ParagraphStyle('h1', parent=styles['Heading1'], fontSize=13,
                                  spaceAfter=8, spaceBefore=14)
        style_h2 = ParagraphStyle('h2', parent=styles['Heading2'], fontSize=11,
                                  spaceAfter=6, spaceBefore=10)
        style_body = ParagraphStyle('body', parent=styles['Normal'], fontSize=9,
                                    spaceAfter=6, leading=13)
        style_bold = ParagraphStyle('bold', parent=styles['Normal'], fontSize=9,
                                    fontName='Helvetica-Bold', spaceAfter=2)

        meta = report_data.get('metadata', {})
        story = []

        # Titel
        story.append(Paragraph('Audit Bevindingenoverzicht', style_title))
        story.append(Paragraph(
            f"Gegenereerd: {datetime.now().strftime('%d-%m-%Y')} &nbsp;|&nbsp; "
            f"Kijktermijn: {meta.get('kijktermijn_jaren', 'N/A')} jaar &nbsp;|&nbsp; "
            f"Bestanden: {meta.get('bestanden_geanalyseerd', 'N/A')}",
            style_sub
        ))
        story.append(PageBreak())

        # Samenvatting
        samenvatting = report_data.get('samenvatting', '')
        if samenvatting:
            story.append(Paragraph('Samenvatting', style_h1))
            story.append(Paragraph(samenvatting, style_body))
            story.append(Spacer(1, 0.5 * cm))

        # Bevindingen
        bevindingen = report_data.get('bevindingen', [])
        if bevindingen:
            story.append(Paragraph('Relevante Bevindingen', style_h1))

            prio_colors = {
                'Hoog': colors.HexColor('#FEE2E2'),
                'Middel': colors.HexColor('#FEF9C3'),
                'Laag': colors.HexColor('#DCFCE7'),
            }

            for prioriteit in ['Hoog', 'Middel', 'Laag']:
                groep = [b for b in bevindingen if b.get('prioriteit') == prioriteit]
                if not groep:
                    continue

                story.append(Paragraph(f'Prioriteit: {prioriteit}', style_h2))

                table_data = [[
                    Paragraph('<b>Bron</b>', style_body),
                    Paragraph('<b>Jaar</b>', style_body),
                    Paragraph('<b>Bevinding</b>', style_body),
                    Paragraph('<b>Relevantie</b>', style_body),
                ]]
                for b in groep:
                    table_data.append([
                        Paragraph(str(b.get('bron', '')), style_body),
                        Paragraph(str(b.get('jaar', '')), style_body),
                        Paragraph(str(b.get('bevinding', '')), style_body),
                        Paragraph(str(b.get('relevantie', '')), style_body),
                    ])

                col_widths = [3.5 * cm, 1.5 * cm, 7 * cm, 5 * cm]
                tbl = Table(table_data, colWidths=col_widths, repeatRows=1)
                bg = prio_colors.get(prioriteit, colors.white)
                tbl.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A5F')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('BACKGROUND', (0, 1), (-1, -1), bg),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [bg, colors.white]),
                ]))
                story.append(tbl)
                story.append(Spacer(1, 0.4 * cm))

        # Aandachtspunten
        aandachtspunten = report_data.get('aandachtspunten', [])
        if aandachtspunten:
            story.append(Paragraph('Aandachtspunten voor de Auditor', style_h1))
            for ap in aandachtspunten:
                story.append(Paragraph(ap.get('onderwerp', ''), style_bold))
                story.append(Paragraph(ap.get('toelichting', ''), style_body))
                story.append(Spacer(1, 0.3 * cm))

        doc.build(story)
